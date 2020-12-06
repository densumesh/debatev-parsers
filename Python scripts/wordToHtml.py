import docx

def converttoHTML(filepath, filelink, types):
    docFile = docx.Document(filepath)
    all_paras = docFile.paragraphs

    arr = []
    for para in all_paras:

        hasNone = False
        hasOther = False

        hasOther2 = False
        for j in para.runs:
            if j.font.size is None:
                hasNone = True
            else :
                hasOther = True
        if (para.style.name == 'Heading 4'):
            arr.append("tag")
        elif (hasNone and hasOther):
            arr.append("card")
        elif (hasNone and hasOther):
            arr.append("cite")
        else:
            arr.append("unsure")


    arr2 = []
    for ii in range (len(arr)):
        if (arr[ii] == 'card' and ii>3):
            
           
            if(arr2[ii-1] ==  "card body"):
                arr2.append("card body")
            else:
                arr2.append("card body")
                arr2[ii-1] = "card cite"
                arr2[ii-2] = "card tag"

                if(arr[ii-3] != 'card'):
                    arr2[ii-3] = "another card tag"
        else:
            arr2.append("nothing")



    def convertPtoHTML(para):
        ret = "<p>"

        for run in para.runs:
            beg = ""
            end = ""
            if run.font.highlight_color is None:
                beg = beg+""
            else:
                beg = beg + "<span style=\"background-color: #FFFF00\">"
                end =  "</span>"+end
            if run.bold == True:
                beg = beg+"<b>"
                end = "</b>"+ end
            elif run.font.size is None and run.underline == True:
                beg = beg+"<u>"
                end = "</u>"+end

            ret = ret+beg+run.text+end

        ret = ret+ "</p>"

        return ret

    i=0
    allHtml = {}
    from docx import Document
    document = Document()
    j=1
    while i < (len(all_paras)):
        if(arr2[i] == "another card tag"):
            paras = all_paras[i]
            tag = paras.text
            i = i+1
            tag = tag+ " "+ all_paras[i].text
            html = "<p><p><b>"+paras.text+"</p></b><p><b>"+all_paras[i].text+"</p></b>"
            i=i+1
            cite = all_paras[i].text;
            html = html+ convertPtoHTML(all_paras[i])
            i = i+1
            html = html + convertPtoHTML(all_paras[i])
            while arr2[i+1] == "card body":
                i = i+1
                html = html + convertPtoHTML(all_paras[i])
            allHtml["card " + str(j)] = [{"tag": tag+ " " +cite, "cardHtml": html, "filepath": filelink, "dtype": types}]
            j =j+1



        elif (arr2[i] == "card tag"):
            paras = all_paras[i]
            tag = paras.text
            i = i + 1
            tag = tag + " " + all_paras[i].text
            html = "<p><p><b>" + paras.text + "</p></b>";
            i = i + 1
            cite = all_paras[i].text;
            html = html + convertPtoHTML(all_paras[i])
            i = i + 1
            html = html + convertPtoHTML(all_paras[i])
            while(arr2[i+1] == "card body"):
                i = i+1
                html = html + convertPtoHTML(all_paras[i])
            allHtml["card " + str(j)] = [{"tag": tag+ " " +cite, "cardHtml": html, "filepath": filelink, 'dtype': types}]
            j = j+1
        i+=1



    return allHtml
